from question import Question
from docx import Document
import lxml
import docx
import sys

def make_q(paragraph):
        parsedq = []
       
        parsedq.extend(paragraph.split('*'))
        for val in parsedq:
                forbidden = ['','/','/n']
                if val in forbidden:
                        parsedq.remove(val)
        
        return parsedq

def parse_q(doc_title):
    questions = dict()
    allq = []
    count = 0
    check = 0
    count1 = 0

    document = docx.Document(doc_title)
    size = len(document.paragraphs)
    questions = dict()
    for x in range(size):
            
            if (document.paragraphs[x].style.name == 'Heading 1'):
                
                q1 = Question('q','m','t')
                count += 1
                question = make_q(document.paragraphs[x].text)
                
                q1.question = question[1]
                q1.metric = question[0]
                q1.type = question[2]
                allq.append(q1)
                check = 1

            elif(document.paragraphs[x].style.name == 'Normal' and document.paragraphs[x].text != "" and check == 1):
                allq[-1].answers.extend(make_q(document.paragraphs[x].text))

            elif(document.paragraphs[x].style.name == 'Normal' and document.paragraphs[x].text == "" ):
                check = 0
                

            elif(document.paragraphs[x].style.name == 'Normal' and document.paragraphs[x].text != ""):
                check = 0
                error = count
                questions[400] = "There was an error codifying your survey, Question number: {}".format(error)


                    
    for question in allq:
        qdict = dict()
        questions[count1] = qdict
        if(question.get_type() == "GRID"):
                qdict['vertical'] = question.get_answers()
                horizontal = question.get_question().split(',')
                horizontal = horizontal[1:]
                qdict["horizontal"] = horizontal
                qdict['type'] = question.get_type()
                qdict['question'] = question.get_question()
                qdict["metric"] = question.get_metric()
        else:
                qdict["metric"] = question.get_metric()
                qdict['question'] = question.get_question()
                qdict['answers'] = question.get_answers()
                qdict['type'] = question.get_type()
        count1+=1

    return questions


